"""Trích văn bản từ PDF và DOCX để nạp vào corpus.

Vì sao cần: giảng viên có sẵn giáo trình dạng PDF/DOCX, bắt họ chuyển tay sang
Markdown thì không ai làm.

Chất lượng hai định dạng KHÔNG ngang nhau, và chênh lệch nằm ở chỗ ít ai ngờ —
tiêu đề mục. Bộ cắt đoạn ưu tiên tách theo ``\\n## `` trước mọi thứ khác (xem
``ingest.split``), nên tài liệu có tiêu đề rõ ràng thì mỗi đoạn ra trọn một ý,
còn tài liệu chỉ toàn văn xuôi thì bị cắt ở mốc 800 ký tự bất kể đang giữa câu
chuyện nào.

    DOCX  giữ style nên khôi phục được: Heading 1/2/3 -> #, ##, ###
    PDF   không có khái niệm tiêu đề, chỉ có chữ đặt ở toạ độ nào đó

Vậy nên **DOCX cho chất lượng truy xuất tốt hơn PDF rõ rệt**. Khuyên giảng viên
nộp DOCX khi có cả hai bản.

Giới hạn đã biết của PDF, không định khắc phục ở đây:

- PDF quét ảnh (scan) trích ra chuỗi rỗng. Cần OCR, nằm ngoài phạm vi — tài
  liệu như vậy bị bỏ qua kèm cảnh báo chứ không nạp bản rỗng.
- Đầu trang, chân trang, số trang lẫn vào nội dung.
- Bố cục nhiều cột dễ bị đọc sai thứ tự.
"""

from __future__ import annotations

from collections.abc import Iterator
from pathlib import Path

import docx
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from ioes_common import get_logger
from pypdf import PdfReader

logger = get_logger(__name__)

#: Đuôi file mà ``ingest.load_corpus`` chấp nhận.
MARKDOWN_SUFFIX = ".md"
SUPPORTED_SUFFIXES = (MARKDOWN_SUFFIX, ".pdf", ".docx")

# Word đặt tên style theo ngôn ngữ giao diện lúc soạn thảo, nên bản tiếng Việt
# cho ra "Đầu đề 1" chứ không phải "Heading 1". Bắt cả hai.
_HEADING_PREFIXES = ("heading ", "đầu đề ", "dau de ")
_TITLE_STYLES = ("title", "tiêu đề", "tieu de")

# Tiêu đề sâu hơn mức này gộp chung vào ###: bộ cắt chỉ tách theo ## và ###,
# thêm mức nữa cũng không đổi được gì.
_MAX_HEADING_LEVEL = 3


def _clean(text: str) -> str:
    """Chuẩn hoá một dòng lấy từ file văn phòng.

    Word chèn tab giữa số thứ tự và nội dung tiêu đề, nên chuỗi lấy ra có dạng
    ``"2.1\\t📑 MỤC LỤC"``. Đo trên PROJECT_MANAGEMENT_PLAN.docx: 103 tab, đúng
    bằng số tiêu đề. Tab lọt vào văn bản đem nhúng chỉ là nhiễu, và tiêu đề còn
    được đem hiện cho người duyệt xem nguồn.

    Chỉ đổi tab thành khoảng trắng và cắt đuôi — KHÔNG gộp khoảng trắng liên
    tiếp, vì nội dung có thể là mã nguồn hoặc bảng cần giữ căn lề.
    """
    return text.replace("\t", " ").strip()


def _heading_level(style_name: str) -> int | None:
    """Mức tiêu đề của một đoạn Word, hoặc None nếu là văn thường."""
    name = (style_name or "").strip().lower()

    if any(name.startswith(prefix) for prefix in _TITLE_STYLES):
        return 1

    for prefix in _HEADING_PREFIXES:
        if name.startswith(prefix):
            tail = name[len(prefix) :].strip()
            if tail.isdigit():
                return min(int(tail), _MAX_HEADING_LEVEL)
    return None


def _table_to_markdown(table: Table) -> str:
    """Đổi bảng Word thành bảng Markdown.

    Bảng hay chứa đúng loại nội dung ra đề được — mã trạng thái HTTP, so sánh
    hai khái niệm. Bỏ qua bảng là mất nội dung mà không báo gì.
    """
    rows = []
    for row in table.rows:
        cells = [_clean(cell.text).replace("|", r"\|").replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append("| " + " | ".join(cells) + " |")

    if not rows:
        return ""

    # Dòng phân cách sau hàng đầu để Markdown nhận ra đây là bảng.
    separator = "| " + " | ".join("---" for _ in table.rows[0].cells) + " |"
    return "\n".join([rows[0], separator, *rows[1:]])


def _iter_block_items(document: DocxDocument) -> Iterator[Paragraph | Table]:
    """Duyệt đoạn văn và bảng theo đúng thứ tự xuất hiện trong tài liệu.

    ``document.paragraphs`` và ``document.tables`` là hai danh sách tách rời,
    nối chúng lại sẽ dồn mọi bảng xuống cuối bài — bảng rời khỏi đoạn văn giải
    thích nó, và đoạn cắt ra mất ngữ cảnh. Phải đọc thẳng cây XML thân bài.
    """
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(child, document)
        elif tag == "tbl":
            yield Table(child, document)


def extract_docx(path: Path) -> tuple[str, str]:
    """Đọc DOCX. Trả về ``(tiêu đề, thân bài dạng Markdown)``.

    Tiêu đề lấy theo thứ tự: thuộc tính tài liệu, rồi tiêu đề mức 1 đầu tiên,
    cuối cùng là tên file.
    """
    document = docx.Document(str(path))

    blocks: list[str] = []
    first_heading = ""

    for item in _iter_block_items(document):
        if isinstance(item, Table):
            markdown_table = _table_to_markdown(item)
            if markdown_table:
                blocks.append(markdown_table)
            continue

        text = _clean(item.text)
        if not text:
            continue

        level = _heading_level(item.style.name if item.style else "")
        if level is None:
            blocks.append(text)
            continue

        if not first_heading:
            first_heading = text
        blocks.append(f"{'#' * level} {text}")

    core_title = _clean(document.core_properties.title or "")
    title = core_title or first_heading or path.stem
    return title, "\n\n".join(blocks)


def extract_pdf(path: Path) -> tuple[str, str]:
    """Đọc PDF. Trả về ``(tiêu đề, thân bài)``.

    Không khôi phục được tiêu đề mục, nên thân bài là văn xuôi thuần — đoạn cắt
    ra sẽ kém mạch lạc hơn DOCX. Xem docstring đầu module.
    """
    reader = PdfReader(str(path))

    pages = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 - một trang hỏng không nên chặn cả file
            logger.warning("pdf_page_unreadable", file=path.name, page=number, error=str(exc))
            continue
        text = _clean(text)
        if text:
            pages.append(text)

    meta_title = ""
    if reader.metadata:
        meta_title = _clean(reader.metadata.title or "")

    return meta_title or path.stem, "\n\n".join(pages)


def extract(path: Path) -> tuple[str, str]:
    """Trích ``(tiêu đề, thân bài)`` theo đuôi file.

    Chuỗi thân bài rỗng nghĩa là không đọc được nội dung — bên gọi phải bỏ qua
    tài liệu đó thay vì nạp một bản rỗng vào vectorstore.
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    raise ValueError(f"Không hỗ trợ đuôi file: {suffix}")
