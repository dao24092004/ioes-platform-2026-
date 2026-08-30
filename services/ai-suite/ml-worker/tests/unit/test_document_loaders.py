"""Kiem thu trich xuat PDF va DOCX.

File thu duoc dung tai cho bang chinh pypdf/python-docx, khong commit file nhi
phan vao repo.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest
from pypdf import PdfWriter

from ml_worker.services import document_loaders, ingest


def _docx(path: Path) -> Path:
    """Mot tai lieu Word co tieu de, van xuoi va mot bang."""
    document = docx.Document()
    document.core_properties.title = "Giáo trình mạng máy tính"

    document.add_heading("Giao thức HTTP", level=1)
    document.add_paragraph("HTTP là giao thức tầng ứng dụng, hoạt động theo mô hình client-server.")

    document.add_heading("Mã trạng thái", level=2)
    document.add_paragraph("Mã trạng thái cho biết kết quả của yêu cầu.")

    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Mã"
    table.cell(0, 1).text = "Ý nghĩa"
    table.cell(1, 0).text = "200"
    table.cell(1, 1).text = "Thành công"
    table.cell(2, 0).text = "404"
    table.cell(2, 1).text = "Không tìm thấy"

    document.save(str(path))
    return path


def _empty_pdf(path: Path, pages: int = 2) -> Path:
    """PDF khong co lop van ban — dung mo phong ban quet anh."""
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    with path.open("wb") as handle:
        writer.write(handle)
    return path


# --- DOCX ---------------------------------------------------------------------


def test_docx_title_comes_from_document_properties(tmp_path: Path) -> None:
    title, _ = document_loaders.extract_docx(_docx(tmp_path / "mang.docx"))

    assert title == "Giáo trình mạng máy tính"


def test_docx_headings_become_markdown(tmp_path: Path) -> None:
    """Bo cat uu tien tach theo '\\n## ', nen tieu de phai thanh Markdown that.

    Mat buoc nay thi ca tai lieu thanh mot khoi van xuoi, bi cat o moc 800 ky tu
    bat ke dang giua y nao.
    """
    _, body = document_loaders.extract_docx(_docx(tmp_path / "mang.docx"))

    assert "# Giao thức HTTP" in body
    assert "\n## Mã trạng thái" in body


def test_docx_keeps_tables(tmp_path: Path) -> None:
    """Bang hay chua dung loai noi dung ra de duoc, bo qua la mat noi dung."""
    _, body = document_loaders.extract_docx(_docx(tmp_path / "mang.docx"))

    assert "| 404 | Không tìm thấy |" in body


def test_docx_keeps_table_next_to_its_paragraph(tmp_path: Path) -> None:
    """Bang phai nam dung cho, khong bi don xuong cuoi bai.

    document.paragraphs va document.tables la hai danh sach roi; noi chung lai
    se tach bang khoi doan van giai thich no.
    """
    _, body = document_loaders.extract_docx(_docx(tmp_path / "mang.docx"))

    assert body.index("Mã trạng thái cho biết") < body.index("| 404 |")


def test_docx_falls_back_to_first_heading(tmp_path: Path) -> None:
    path = tmp_path / "khong-tieu-de.docx"
    document = docx.Document()
    document.add_heading("Kiến trúc microservice", level=1)
    document.add_paragraph("Mỗi service là một bounded context.")
    document.save(str(path))

    title, _ = document_loaders.extract_docx(path)

    assert title == "Kiến trúc microservice"


# --- PDF ----------------------------------------------------------------------


def test_scanned_pdf_yields_empty_body(tmp_path: Path) -> None:
    """PDF quet anh khong co lop van ban nen trich ra chuoi rong."""
    _, body = document_loaders.extract_pdf(_empty_pdf(tmp_path / "scan.pdf"))

    assert body == ""


def test_pdf_title_falls_back_to_filename(tmp_path: Path) -> None:
    title, _ = document_loaders.extract_pdf(_empty_pdf(tmp_path / "bai-giang-01.pdf"))

    assert title == "bai-giang-01"


# --- Ket noi voi load_corpus --------------------------------------------------


def test_load_corpus_reads_docx(tmp_path: Path) -> None:
    _docx(tmp_path / "mang.docx")

    documents = ingest.load_corpus(tmp_path)

    assert len(documents) == 1
    assert documents[0].metadata["title"] == "Giáo trình mạng máy tính"
    # doc_id lay ten file, nen chunk_id trich dan se la mang#0, mang#1...
    assert documents[0].metadata["doc_id"] == "mang"


def test_load_corpus_skips_document_with_no_text(tmp_path: Path) -> None:
    """Nap ban rong thi no van chiem cho trong ket qua truy xuat ma khong co gi de doc."""
    _empty_pdf(tmp_path / "scan.pdf")
    _docx(tmp_path / "mang.docx")

    documents = ingest.load_corpus(tmp_path)

    assert [d.metadata["doc_id"] for d in documents] == ["mang"]


def test_load_corpus_skips_unreadable_file_without_failing_the_run(
    tmp_path: Path,
) -> None:
    """Mot file hong khong duoc chan ca lan nap."""
    (tmp_path / "hong.docx").write_bytes(b"day khong phai file docx")
    _docx(tmp_path / "mang.docx")

    documents = ingest.load_corpus(tmp_path)

    assert [d.metadata["doc_id"] for d in documents] == ["mang"]


def test_load_corpus_ignores_unsupported_suffixes(tmp_path: Path) -> None:
    (tmp_path / "anh.png").write_bytes(b"\x89PNG")
    (tmp_path / "bang.xlsx").write_bytes(b"PK")
    _docx(tmp_path / "mang.docx")

    documents = ingest.load_corpus(tmp_path)

    assert len(documents) == 1


def test_load_corpus_still_reads_markdown_frontmatter(tmp_path: Path) -> None:
    (tmp_path / "vi-git.md").write_text(
        "---\ntitle: Git và GitHub\ndoc_id: vi-git\n---\n\nRebase viết lại lịch sử.",
        encoding="utf-8",
    )

    documents = ingest.load_corpus(tmp_path)

    assert documents[0].metadata == {
        "doc_id": "vi-git",
        "title": "Git và GitHub",
        "lang": "en",
    }


def test_extract_rejects_unsupported_suffix(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="Không hỗ trợ"):
        document_loaders.extract(tmp_path / "bang.xlsx")


def test_strips_word_numbering_tabs_from_headings(tmp_path: Path) -> None:
    """Word chen tab giua so thu tu va noi dung tieu de.

    Do tren PROJECT_MANAGEMENT_PLAN.docx: 103 tab, dung bang so tieu de. Tab lot
    vao van ban dem nhung chi la nhieu, va tieu de con duoc hien cho nguoi duyet
    xem nguon.
    """
    path = tmp_path / "co-tab.docx"
    document = docx.Document()
    document.add_heading("2.1\tMục lục", level=2)
    document.add_paragraph("Nội dung\tcó tab.")
    document.save(str(path))

    title, body = document_loaders.extract_docx(path)

    assert "\t" not in body
    assert "\t" not in title
    assert "## 2.1 Mục lục" in body
